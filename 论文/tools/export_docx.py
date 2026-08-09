"""Export the manuscript Markdown to a review-oriented DOCX.

The Markdown file remains the source of truth.  This exporter intentionally
implements only the structures used by the manuscript: headings, paragraphs,
pipe tables, the architecture figure, and simple inline emphasis/code.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "GriddingMachine论文初稿_v0.1.md"
OUTPUT = ROOT / "GriddingMachine论文初稿_导师审阅版.docx"
FIGURE_PREVIEW = ROOT / "figures" / "figure1-preview.png"


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman") -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        else:
            run = paragraph.add_run(part)
        if not (part.startswith("`") and part.endswith("`")):
            set_run_font(run)


def table_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator(line: str) -> bool:
    cells = table_cells(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c) for c in cells)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0.55)
    paragraph.paragraph_format.right_indent = Cm(0.55)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.0
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F5F7")
    p_pr.append(shd)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")


def add_table(document: Document, lines: list[str]) -> None:
    rows = [table_cells(line) for line in lines if not is_separator(line)]
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(width):
            text = row[j] if j < len(row) else ""
            cell = table.cell(i, j)
            cell.text = ""
            add_inline(cell.paragraphs[0], text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True
            if i == 0:
                shade_cell(cell, "D9EAF7")


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.space_after = Pt(0)

    for name, size in (("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11)):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.paragraph_format.first_line_indent = Cm(0)


def export() -> None:
    document = Document()
    configure(document)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith("!["):
            if FIGURE_PREVIEW.exists():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(FIGURE_PREVIEW), width=Cm(15.0))
            i += 1
            continue

        if line.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1 if i < len(lines) else 0
            add_code_block(document, code_lines)
            continue

        if line.startswith("|") and i + 1 < len(lines) and is_separator(lines[i + 1].strip()):
            block = [line, lines[i + 1].strip()]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            add_table(document, block)
            continue

        match = re.match(r"^(#{1,4})\s+(.*)$", line)
        if match:
            level = len(match.group(1))
            title = match.group(2)
            if level == 1:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(paragraph, title)
            else:
                paragraph = document.add_heading(level=min(level - 1, 3))
                add_inline(paragraph, title)
            i += 1
            continue

        paragraph = document.add_paragraph()
        if line.startswith("- "):
            paragraph.style = "List Bullet"
            line = line[2:]
        add_inline(paragraph, line)
        if line.startswith("**表") or line.startswith("**Table") or line.startswith("**图") or line.startswith("**Fig."):
            paragraph.paragraph_format.first_line_indent = Cm(0)
        i += 1

    core = document.core_properties
    core.title = "GriddingMachine：全球网格数据生产、分发与模型调用框架的更新与验证"
    core.author = "Hao Jiang; Yujie Wang"
    core.subject = "导师审阅版，由Markdown源稿自动生成"
    document.add_section(WD_SECTION.NEW_PAGE)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    export()
